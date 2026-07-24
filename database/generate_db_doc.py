#!/usr/bin/env python3
"""
Menghasilkan dokumentasi struktur basis data aplikasi (dalam Bahasa Indonesia)
sebagai file .docx, berdasarkan metadata dari database yang sedang berjalan
(information_schema) ditambah anotasi manual (fungsi tabel/kolom, status
pemakaian) yang sudah diverifikasi terhadap kode aplikasi.

Penggunaan: python3 database/generate_db_doc.py [output.docx]
"""
import os
import sys

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from generate_erd import load_env, fetch_schema, merge_manual_tables  # noqa: E402
from schema_modules import MODULES  # noqa: E402

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

NAVY = RGBColor(0x1E, 0x29, 0x3B)
MUTED = RGBColor(0x47, 0x55, 0x69)
ACCENT = RGBColor(0x25, 0x63, 0xEB)

# ---------------------------------------------------------------------------
# Anotasi manual: fungsi tabel, keterangan per kolom, catatan status pemakaian.
# Diverifikasi terhadap src/Controllers & src/Repositories, bukan tebakan dari
# nama kolom saja.
# ---------------------------------------------------------------------------

TABLE_DOCS = {
    'users': {
        'fungsi': 'Tabel utama akun pengguna aplikasi untuk ketiga peran (admin, konselor, mahasiswa). Menyimpan data login sekaligus data profil dasar; data profil khusus konselor disimpan terpisah di tabel counselors.',
        'kolom': {
            'id': 'Kunci utama (PK), dipakai sebagai acuan oleh hampir seluruh tabel lain lewat kolom user_id.',
            'name': 'Nama panggilan/tampilan pengguna.',
            'username': 'Username untuk login.',
            'password': 'Hash password (bcrypt, via password_hash PHP) — tidak pernah disimpan dalam bentuk teks biasa.',
            'full_name': 'Nama lengkap. Hanya diisi mahasiswa yang mendaftar mandiri lewat halaman Daftar; untuk admin/konselor yang akunnya dibuat oleh admin, kolom ini kosong dan kolom name berperan sebagai nama lengkap.',
            'student_number': 'NPM, khusus untuk akun mahasiswa.',
            'gender': "Jenis kelamin ('Male'/'Female').",
            'faculty': 'Nama fakultas, khusus untuk akun mahasiswa.',
            'major': 'Nama jurusan, khusus untuk akun mahasiswa.',
            'phone_number': 'Nomor HP.',
            'email': 'Alamat email — dipakai juga untuk mengirim tautan reset password dan notifikasi persetujuan akun.',
            'role': "Peran akun: 'admin', 'student' (mahasiswa), atau 'counselor' (konselor).",
            'profile_image': 'Nama file foto profil yang diunggah lewat halaman Profil (disimpan di public/uploads/profile).',
            'status': "Status akun: 'pending' (baru daftar, menunggu persetujuan admin), 'active' (bisa login), atau 'rejected' (pendaftaran ditolak).",
            'approved_by': 'ID admin yang menyetujui/menolak pendaftaran akun ini — referensi ke users.id itu sendiri (self-reference).',
            'approved_at': 'Waktu akun disetujui/ditolak.',
            'created_at': 'Waktu akun dibuat.',
        },
        'relasi_extra': [
            'approved_by merupakan referensi ke tabel users itu sendiri (satu admin bisa menyetujui banyak akun lain).',
            'Menjadi tabel induk (parent) bagi hampir seluruh tabel lain di aplikasi melalui kolom user_id/counselor_id, karena baik data mahasiswa maupun data login konselor sama-sama berakar dari tabel ini.',
        ],
    },
    'password_reset_tokens': {
        'fungsi': 'Menyimpan token sekali-pakai untuk fitur lupa password. Hanya hash SHA-256 dari token yang disimpan — token asli hanya dikirim lewat email dan tidak pernah disimpan, agar kebocoran data tidak bisa langsung dipakai mereset password siapa pun.',
        'kolom': {
            'id': 'Kunci utama (PK).',
            'user_id': 'Pemilik permintaan reset password.',
            'token_hash': 'Hash SHA-256 dari token yang dikirim lewat email (kolom unik).',
            'expires_at': 'Batas waktu token berlaku (60 menit sejak dibuat).',
            'used_at': 'Waktu token dipakai untuk mengganti password; kosong berarti token belum/tidak pernah dipakai.',
            'created_at': 'Waktu permintaan reset dibuat.',
        },
    },
    'faculties': {
        'fungsi': 'Data referensi daftar fakultas, dipakai saat pendaftaran akun mahasiswa (dropdown pilihan Fakultas).',
        'kolom': {
            'id': 'Kunci utama (PK).',
            'name': 'Nama fakultas.',
        },
    },
    'majors': {
        'fungsi': 'Data referensi daftar jurusan per fakultas, dipakai saat pendaftaran akun mahasiswa (dropdown Jurusan yang berubah mengikuti Fakultas yang dipilih).',
        'kolom': {
            'id': 'Kunci utama (PK).',
            'faculty_id': 'Fakultas tempat jurusan ini berada.',
            'name': 'Nama jurusan.',
        },
    },
    'counselors': {
        'fungsi': 'Data profil konsultasi milik akun berperan konselor — perpanjangan (1-ke-1) dari sebuah baris di tabel users. Dibuat/diubah lewat menu admin "Kelola Konselor".',
        'kolom': {
            'counselor_id': 'Kunci utama (PK) — dipakai sebagai acuan oleh booking, jadwal, rating, dsb (bukan user_id, karena satu konselor = satu counselor_id yang tetap meski data user berubah).',
            'user_id': 'Akun login (tabel users) pemilik profil konselor ini.',
            'registration_number': 'Nomor registrasi/izin praktik (NIP/NIK), harus unik.',
            'profession': "Profesi: 'Psychologist' (Psikolog), 'Counselor' (Konselor), atau 'Psychiatrist' (Psikiater).",
            'specialization': 'Bidang spesialisasi (mis. Konseling Akademik).',
            'education': 'Riwayat pendidikan (mis. S2 Psikologi).',
            'experience_years': 'Lama pengalaman praktik, dalam tahun.',
            'languages': 'Bahasa yang dikuasai.',
            'consultation_fee': 'Biaya konsultasi (kolom tersimpan, namun saat ini disembunyikan/di-comment di tampilan direktori konselor).',
            'session_duration': 'Durasi standar satu sesi konsultasi, dalam menit.',
            'consultation_method': "Metode konsultasi: 'Online', 'Offline', atau 'Hybrid'.",
            'profile_photo': 'Path lengkap foto profil konsultasi (mis. /uploads/profile/xxx.jpg) — diisi lewat form admin, terpisah dari users.profile_image yang diisi mahasiswa/konselor sendiri lewat halaman Profil.',
            'biography': 'Biografi/deskripsi singkat yang tampil di halaman profil konselor.',
            'verification_status': 'Penanda status verifikasi profil oleh admin (1/0).',
            'is_active': 'Menentukan apakah konselor tampil di direktori konselor dan bisa menerima booking (1) atau tidak (0).',
            'created_at': 'Waktu profil dibuat.',
            'updated_at': 'Waktu profil terakhir diubah.',
        },
        'relasi_extra': [
            'user_id -> users.id (relasi satu-ke-satu: satu akun konselor punya tepat satu profil konselor).',
        ],
    },
    'counselor_schedules': {
        'fungsi': 'Slot jadwal konsultasi yang dibuka oleh admin untuk seorang konselor. Mahasiswa memilih salah satu slot ini saat mengajukan booking.',
        'kolom': {
            'schedule_id': 'Kunci utama (PK).',
            'counselor_id': 'Konselor pemilik slot jadwal ini.',
            'date': 'Tanggal slot.',
            'start_time': 'Jam mulai.',
            'end_time': 'Jam selesai.',
            'quota': 'Kuota maksimal booking yang bisa menempati slot ini.',
            'is_active': 'Menentukan apakah slot ini masih bisa dipilih mahasiswa (1) atau dinonaktifkan (0).',
            'created_at': 'Waktu slot dibuat.',
        },
    },
    'counseling_bookings': {
        'fungsi': 'Pengajuan booking konsultasi dari mahasiswa ke konselor untuk satu slot jadwal tertentu. Statusnya berubah seiring alur konfirmasi konselor dan (sejak fitur persetujuan pembatalan) persetujuan admin.',
        'kolom': {
            'booking_id': 'Kunci utama (PK).',
            'user_id': 'Mahasiswa yang mengajukan booking.',
            'counselor_id': 'Konselor tujuan booking.',
            'schedule_id': 'Slot jadwal (counselor_schedules) yang dipilih.',
            'date': 'Tanggal konsultasi (disalin dari slot jadwal saat booking dibuat).',
            'start_time': 'Jam mulai konsultasi.',
            'end_time': 'Jam selesai konsultasi.',
            'complaint': 'Keluhan singkat yang dituliskan mahasiswa saat mengajukan booking (opsional).',
            'status': "Status booking: 'Pending' (menunggu konfirmasi konselor), 'Confirmed' (dikonfirmasi, monitoring aktif), 'Completed' (sesi selesai), 'Cancelled' (dibatalkan), 'No Show' (mahasiswa tidak hadir), atau 'Cancellation Requested' (mahasiswa mengajukan pembatalan, menunggu keputusan admin).",
            'created_at': 'Waktu booking diajukan.',
        },
        'relasi_extra': [
            'Menjadi acuan utama bagi booking_cancellation_requests, monitoring_periods, counseling_sessions, dan counselor_ratings.',
        ],
    },
    'booking_cancellation_requests': {
        'fungsi': (
            'Menyimpan pengajuan pembatalan booking oleh mahasiswa yang menunggu persetujuan admin. '
            'Ditambahkan agar pembatalan tidak langsung terjadi begitu mahasiswa menekan tombol "Batal" — booking '
            'diparkir berstatus "Cancellation Requested" sampai admin menyetujui (booking menjadi "Cancelled") atau '
            'menolak (booking kembali ke status semula lewat kolom previous_status).'
        ),
        'kolom': {
            'id': 'Kunci utama (PK).',
            'booking_id': 'Booking yang diajukan pembatalannya.',
            'previous_status': "Status booking sebelum pembatalan diajukan ('Pending' atau 'Confirmed') — dipakai untuk mengembalikan status booking jika admin menolak permintaan.",
            'reason': 'Alasan pembatalan yang dituliskan mahasiswa (opsional).',
            'status': "Status pengajuan pembatalan itu sendiri: 'Pending' (menunggu admin), 'Approved' (disetujui), atau 'Rejected' (ditolak).",
            'admin_notes': 'Catatan admin saat menolak pengajuan (opsional).',
            'reviewed_by': 'Admin yang memproses (menyetujui/menolak) pengajuan ini.',
            'reviewed_at': 'Waktu pengajuan diproses admin.',
            'created_at': 'Waktu mahasiswa mengajukan pembatalan.',
        },
    },
    'monitoring_periods': {
        'fungsi': (
            'Jendela waktu pemantauan yang aktif setelah konselor mengonfirmasi sebuah booking. Selama periode ini '
            'berlangsung, mahasiswa & konselor bisa saling chat dan mahasiswa bisa membagikan Diary ke konselor '
            'tersebut. Satu booking hanya punya satu periode monitoring (kolom booking_id bersifat unik).'
        ),
        'kolom': {
            'monitoring_id': 'Kunci utama (PK).',
            'booking_id': 'Booking (yang sudah Confirmed) yang memulai periode monitoring ini.',
            'user_id': 'Mahasiswa yang dipantau.',
            'counselor_id': 'Konselor yang memantau.',
            'start_date': 'Tanggal mulai periode monitoring.',
            'end_date': 'Tanggal berakhir periode monitoring (bisa diperpanjang konselor, atau diakhiri lebih awal saat booking Completed/No Show/pembatalan disetujui).',
            'created_at': 'Waktu periode monitoring dibuat.',
        },
    },
    'counseling_sessions': {
        'fungsi': 'Catatan hasil sesi konseling yang diisi konselor saat menandai sebuah booking selesai (fitur "Tandai Selesai"). Satu booking hanya punya satu catatan sesi (booking_id unik).',
        'kolom': {
            'session_id': 'Kunci utama (PK).',
            'booking_id': 'Booking yang sesinya dicatat di sini.',
            'counselor_notes': 'Ringkasan/catatan konselor tentang sesi konseling.',
            'recommendation': 'Rekomendasi konselor untuk mahasiswa.',
            'follow_up': 'Rencana tindak lanjut.',
            'duration': 'Lama sesi berlangsung (menit) — kolom tersimpan namun tidak diisi oleh form saat ini.',
            'completed_at': 'Waktu sesi ditandai selesai — kolom tersimpan namun tidak diisi oleh form saat ini (memakai created_at).',
            'created_at': 'Waktu catatan sesi ini dibuat/diperbarui.',
        },
        'relasi_extra': [
            'Data pada tabel ini yang menjadi sumber utama Laporan Konseling (Laporan > Konseling).',
        ],
    },
    'chat_messages': {
        'fungsi': 'Riwayat percakapan chat antara mahasiswa dan konselor, hanya bisa terjadi selama monitoring_periods sedang aktif untuk pasangan mahasiswa-konselor tersebut.',
        'kolom': {
            'id': 'Kunci utama (PK).',
            'user_id': 'Pengirim pesan (nama kolom historis; secara aplikasi kolom ini berfungsi sebagai ID pengirim, bukan "pemilik percakapan").',
            'sender_id': 'Kolom lama yang sudah tidak dipakai kode aplikasi saat ini (pengiriman pesan hanya mengisi user_id dan receiver_id).',
            'receiver_id': 'Penerima pesan.',
            'message': 'Isi pesan.',
            'created_at': 'Waktu pesan dikirim.',
            'is_read': 'Penanda apakah pesan sudah dibaca penerima (dipakai untuk badge jumlah pesan belum dibaca di kotak masuk konselor).',
        },
    },
    'diary_entries': {
        'fungsi': (
            'Diary Terstruktur — catatan harian mahasiswa berbasis pendekatan CBT (Cognitive Behavioral Therapy): '
            'situasi, pikiran otomatis, emosi, reaksi fisik, perilaku, refleksi diri, dan rasa syukur. Bisa diset '
            'privat atau dibagikan ke satu konselor tertentu.'
        ),
        'kolom': {
            'id': 'Kunci utama (PK).',
            'user_id': 'Mahasiswa pemilik catatan diary.',
            'entry_date': 'Tanggal catatan diary.',
            'title': 'Kolom judul dari versi diary bebas (free-text) sebelumnya — tidak lagi diisi/dipakai oleh fitur Diary Terstruktur saat ini.',
            'content': 'Kolom isi dari versi diary bebas (free-text) sebelumnya — tidak lagi diisi/dipakai oleh fitur Diary Terstruktur saat ini.',
            'situation': 'Bagian 1: situasi/kejadian yang membuat tidak nyaman.',
            'initial_thoughts': 'Bagian 2: pikiran pertama/otomatis yang muncul.',
            'emotions_list': 'Bagian 3: daftar emosi yang dirasakan (disimpan sebagai JSON array).',
            'other_emotions': 'Emosi lain di luar pilihan baku, jika mahasiswa memilih opsi "Lainnya".',
            'emotion_intensity': 'Skala intensitas emosi, 1 (sangat ringan) sampai 5 (sangat berat).',
            'physical_reactions_list': 'Bagian 4: daftar reaksi fisik yang dirasakan (JSON array).',
            'other_physical_reactions': 'Reaksi fisik lain di luar pilihan baku.',
            'behavior': 'Bagian 5: perilaku yang dilakukan setelah kejadian.',
            'self_reflection': 'Refleksi diri (opsional) — juga ditampilkan di fitur Self Help > Syukur & Refleksi Diri.',
            'gratitude_list': 'Jurnal syukur, 3 hal yang disyukuri hari itu (JSON array) — juga ditampilkan di fitur Self Help.',
            'tomorrow_plan': 'Rencana kecil untuk besok (opsional).',
            'mood_level': 'Kolom mood dari versi diary sebelumnya — tidak lagi diisi/dipakai oleh fitur Diary Terstruktur saat ini.',
            'is_private': 'Menentukan apakah catatan ini privat (1, hanya mahasiswa yang bisa lihat) atau dibagikan ke konselor (0).',
            'shared_counselor_id': 'Konselor tujuan berbagi, jika is_private = 0.',
            'created_at': 'Waktu catatan dibuat.',
        },
        'relasi_extra': [
            'shared_counselor_id -> counselors.counselor_id (bukan ke users, karena yang dirujuk adalah profil konselornya).',
        ],
    },
    'assessment_questions': {
        'fungsi': 'Bank soal untuk dua instrumen self-assessment yang dipakai aplikasi: BDI-II (Beck Depression Inventory-II, 21 butir) dan PWB (Psychological Well-Being, 18 butir).',
        'kolom': {
            'id': 'Kunci utama (PK).',
            'type': "Instrumen soal ini: 'bdi2' atau 'pwb'.",
            'order_no': 'Nomor urut tampil soal dalam instrumennya.',
            'question_text': 'Teks pertanyaan.',
            'dimension': 'Nama dimensi PWB terkait (mis. Otonomi, Penguasaan Lingkungan) — kosong untuk soal BDI-II.',
            'is_reverse_scored': 'Penanda soal dengan penilaian terbalik (skor tinggi = kondisi kurang baik, khusus sebagian butir PWB).',
            'created_at': 'Waktu soal dibuat.',
        },
    },
    'assessment_choices': {
        'fungsi': 'Pilihan jawaban (skala Likert/pilihan berjenjang) untuk setiap soal di assessment_questions, beserta bobot skornya masing-masing.',
        'kolom': {
            'id': 'Kunci utama (PK).',
            'question_id': 'Soal pemilik pilihan jawaban ini.',
            'order_no': 'Urutan tampil pilihan.',
            'label': 'Teks pilihan jawaban.',
            'score_value': 'Skor yang didapat jika pilihan ini dipilih.',
        },
    },
    'assessment_sessions': {
        'fungsi': (
            'Satu sesi pengisian gabungan BDI-II + PWB yang dibatasi waktu (timer berjalan di server). Mahasiswa '
            'mengisi kedua instrumen berturut-turut dalam satu sesi; jika waktu habis, jawaban yang sempat terisi '
            'otomatis dikirim.'
        ),
        'kolom': {
            'id': 'Kunci utama (PK).',
            'user_id': 'Mahasiswa pemilik sesi.',
            'status': "Status sesi: 'in_progress', 'completed', atau 'timed_out'.",
            'time_limit_seconds': 'Batas waktu pengisian, dalam detik (diambil dari pengaturan Admin > Pengaturan Sistem saat sesi dibuat).',
            'started_at': 'Waktu sesi dimulai.',
            'expires_at': 'Waktu sesi akan otomatis berakhir.',
            'last_seen_question_id': 'Soal terakhir yang sedang dilihat mahasiswa — dipakai agar tampilan bisa melanjutkan dari posisi terakhir jika halaman dimuat ulang.',
            'bdi2_submission_id': 'Hasil akhir BDI-II dari sesi ini, setelah difinalisasi (lihat assessment_submissions).',
            'pwb_submission_id': 'Hasil akhir PWB dari sesi ini, setelah difinalisasi.',
            'finalized_at': 'Waktu sesi difinalisasi (selesai normal maupun karena waktu habis).',
            'created_at': 'Waktu sesi dibuat.',
        },
    },
    'assessment_session_answers': {
        'fungsi': 'Draft jawaban yang tersimpan selama sesi pengisian gabungan (assessment_sessions) masih berlangsung — memungkinkan mahasiswa berpindah antar soal secara bebas sebelum sesi difinalisasi.',
        'kolom': {
            'id': 'Kunci utama (PK).',
            'session_id': 'Sesi pengisian pemilik jawaban ini.',
            'question_id': 'Soal yang dijawab.',
            'choice_id': 'Pilihan jawaban yang dipilih.',
            'score_value': 'Skor dari pilihan yang dipilih (disalin agar perhitungan tidak perlu join berulang).',
            'answered_at': 'Waktu jawaban ini diisi/diubah terakhir.',
        },
    },
    'assessment_submissions': {
        'fungsi': 'Hasil akhir satu instrumen (BDI-II atau PWB) setelah sesi difinalisasi — inilah yang dipakai untuk ditampilkan sebagai "Hasil Assessment" dan dipakai laporan.',
        'kolom': {
            'id': 'Kunci utama (PK).',
            'user_id': 'Mahasiswa pemilik hasil.',
            'type': "Instrumen: 'bdi2' atau 'pwb'.",
            'total_score': 'Total skor akhir.',
            'max_score': 'Skor maksimal instrumen ini (untuk menghitung persentase/kategori).',
            'category': "Kategori hasil, mis. 'Minimal'/'Ringan'/'Sedang'/'Berat' (BDI-II) atau 'Tinggi'/'Sedang'/'Rendah' (PWB).",
            'category_percentage': 'Persentase skor terhadap skor maksimal.',
            'is_timed_out': 'Penanda apakah hasil ini terbentuk karena sesi kehabisan waktu (1) atau selesai normal (0).',
            'dimension_scores': 'Rincian skor per dimensi PWB (JSON) — kosong untuk BDI-II.',
            'submitted_at': 'Waktu hasil ini difinalisasi.',
        },
        'relasi_extra': [
            'Direferensikan balik oleh assessment_sessions.bdi2_submission_id dan pwb_submission_id.',
        ],
    },
    'assessment_answers': {
        'fungsi': 'Salinan permanen jawaban per soal dari sebuah hasil akhir (assessment_submissions) — dipakai untuk menampilkan rincian jawaban di halaman Hasil Assessment.',
        'kolom': {
            'id': 'Kunci utama (PK).',
            'submission_id': 'Hasil akhir (assessment_submissions) pemilik rincian jawaban ini.',
            'question_id': 'Soal yang dijawab.',
            'choice_id': 'Pilihan jawaban yang dipilih.',
            'score_value': 'Skor dari pilihan yang dipilih.',
        },
    },
    'assessment_retake_grants': {
        'fungsi': (
            'Izin mengisi ulang self-assessment. Setelah pengisian pertama, mahasiswa hanya boleh mengisi lagi jika '
            'seorang konselor merekomendasikannya saat menandai sebuah booking selesai (mencentang "Rekomendasikan '
            'Assessment Ulang"). Satu izin hanya berlaku untuk satu kali pengisian ulang.'
        ),
        'kolom': {
            'id': 'Kunci utama (PK).',
            'user_id': 'Mahasiswa penerima izin.',
            'booking_id': 'Booking konseling yang menjadi dasar pemberian izin.',
            'counselor_id': 'Konselor yang memberi rekomendasi.',
            'granted_at': 'Waktu izin diberikan.',
            'consumed_at': 'Waktu izin dipakai (kosong berarti izin belum dipakai/masih berlaku).',
            'consumed_session_id': 'Sesi assessment yang memakai izin ini.',
        },
    },
    'self_help_activities': {
        'fungsi': 'Rencana aktivitas positif kecil yang dicatat mahasiswa lewat fitur Self Help, lengkap dengan perbandingan suasana hati sebelum & sesudah menjalankannya.',
        'kolom': {
            'id': 'Kunci utama (PK).',
            'user_id': 'Mahasiswa pemilik aktivitas.',
            'title': 'Nama aktivitas.',
            'description': 'Catatan tambahan tentang aktivitas (opsional).',
            'planned_date': 'Tanggal aktivitas direncanakan.',
            'mood_before': 'Skor mood sebelum aktivitas (1-5, opsional).',
            'mood_after': 'Skor mood sesudah aktivitas (1-5, diisi saat aktivitas ditandai selesai).',
            'status': "Status: 'planned' (direncanakan), 'done' (selesai), atau 'skipped' (dilewati).",
            'created_at': 'Waktu aktivitas dicatat.',
            'updated_at': 'Waktu terakhir diubah.',
        },
    },
    'articles': {
        'fungsi': 'Artikel edukasi seputar kesehatan mental yang ditulis admin/konselor dan bisa dibaca semua pengguna.',
        'kolom': {
            'id': 'Kunci utama (PK).',
            'admin_id': 'Kolom lama (legacy) — selalu diisi nilai yang sama dengan user_id oleh kode aplikasi saat ini, tidak dipakai terpisah.',
            'title': 'Judul artikel.',
            'content': 'Isi artikel.',
            'category': 'Kategori artikel (opsional).',
            'image': 'Nama file gambar sampul (opsional).',
            'tags': 'Tag artikel, dipisah koma (opsional).',
            'published_at': 'Waktu publikasi.',
            'created_at': 'Waktu artikel dibuat.',
            'user_id': 'Penulis artikel (admin atau konselor) — kolom yang aktif dipakai aplikasi untuk menentukan kepemilikan (tombol Edit/Hapus).',
        },
    },
    'daily_tips': {
        'fungsi': 'Tips harian yang tampil sebagai popup untuk mahasiswa setiap kali login, dikelola oleh konselor.',
        'kolom': {
            'id': 'Kunci utama (PK).',
            'content': 'Isi tips.',
            'is_active': 'Menentukan apakah tips ini termasuk yang aktif ditampilkan (1) atau tidak (0).',
            'created_by': 'Konselor/admin yang membuat tips ini.',
            'created_at': 'Waktu dibuat.',
            'updated_at': 'Waktu terakhir diubah.',
            'title': 'Judul tips.',
        },
    },
    'app_settings': {
        'fungsi': (
            'Pengaturan aplikasi berbentuk pasangan key-value yang bisa diubah admin lewat menu Pengaturan Sistem. '
            "Saat ini baru dipakai untuk satu pengaturan: assessment_time_limit_minutes (batas waktu pengisian "
            'self-assessment, default 45 menit).'
        ),
        'kolom': {
            'setting_key': 'Kunci utama (PK) — nama pengaturan, mis. assessment_time_limit_minutes.',
            'setting_value': 'Nilai pengaturan (disimpan sebagai teks).',
            'updated_at': 'Waktu terakhir diubah.',
        },
    },
    # --- Tabel warisan / tidak digunakan aktif oleh kode aplikasi saat ini ---
    'students': {
        'fungsi': 'Tabel data mahasiswa dari iterasi awal skema database, sebelum data mahasiswa disatukan langsung ke tabel users (kolom student_number/faculty/major/gender di users).',
        'kolom': {
            'student_id': 'Kunci utama (PK).',
            'id': 'Referensi ke users.id.',
            'student_number': 'NPM.',
            'study_program': 'Program studi.',
            'faculty': 'Fakultas.',
            'gender': "Jenis kelamin ('M'/'F') — memakai kode berbeda dari users.gender ('Male'/'Female').",
            'phone_number': 'Nomor HP.',
        },
        'catatan': 'Tabel ini tidak lagi dibaca maupun ditulis oleh kode aplikasi saat ini — seluruh fitur (pendaftaran, daftar mahasiswa, laporan) memakai tabel users langsung.',
    },
    'system_settings': {
        'fungsi': 'Tabel pengaturan generik dari iterasi awal, sebelum digantikan oleh app_settings.',
        'kolom': {
            'id': 'Kunci utama (PK).',
            'setting_key': 'Nama pengaturan.',
            'setting_value': 'Nilai pengaturan.',
        },
        'catatan': 'Tidak digunakan oleh kode aplikasi saat ini — sudah digantikan oleh app_settings.',
    },
    'questions': {
        'fungsi': 'Bank soal self-assessment sederhana dari versi awal aplikasi, sebelum digantikan sistem BDI-II/PWB (assessment_questions dkk).',
        'kolom': {
            'id': 'Kunci utama (PK).',
            'question': 'Teks pertanyaan.',
            'category': 'Kategori soal.',
            'score_weight': 'Bobot skor.',
            'choice_type': 'Tipe pilihan jawaban.',
            'created_at': 'Waktu dibuat.',
        },
        'catatan': 'Tidak digunakan oleh kode aplikasi saat ini — sudah digantikan oleh assessment_questions/assessment_choices.',
    },
    'answers': {
        'fungsi': 'Jawaban dari sistem self-assessment sederhana versi awal (lihat tabel questions).',
        'kolom': {
            'id': 'Kunci utama (PK).',
            'user_id': 'Pengguna yang menjawab.',
            'question_id': 'Soal yang dijawab.',
            'answer_value': 'Nilai jawaban.',
        },
        'catatan': 'Tidak digunakan oleh kode aplikasi saat ini — sudah digantikan oleh assessment_answers/assessment_session_answers.',
    },
    'assessment_results': {
        'fungsi': 'Rangkuman hasil assessment dari versi awal sebelum alur sesi gabungan BDI-II+PWB (assessment_sessions/assessment_submissions) dibuat.',
        'kolom': {
            'id': 'Kunci utama (PK).',
            'user_id': 'Mahasiswa pemilik hasil.',
            'username': 'Salinan username (denormalisasi) dari versi awal.',
            'result_summary': 'Ringkasan hasil.',
            'assessment_date': 'Tanggal assessment.',
            'created_at': 'Waktu dibuat.',
            'test_date': 'Waktu tes dilakukan.',
            'total_score': 'Total skor.',
            'conclusion': 'Kesimpulan hasil.',
            'recommended_action': 'Saran tindak lanjut.',
            'review_status': "Status peninjauan: 'Not Reviewed' atau 'Reviewed'.",
        },
        'catatan': 'Tidak digunakan oleh kode aplikasi saat ini — sudah digantikan oleh assessment_submissions.',
    },
    'counselor_ratings': {
        'fungsi': 'Skema untuk fitur penilaian/ulasan mahasiswa terhadap konselor setelah sesi konseling.',
        'kolom': {
            'rating_id': 'Kunci utama (PK).',
            'booking_id': 'Booking yang dinilai (satu booking hanya bisa dinilai sekali).',
            'user_id': 'Mahasiswa pemberi nilai.',
            'counselor_id': 'Konselor yang dinilai.',
            'rating': 'Nilai/skor rating.',
            'comment': 'Komentar/ulasan.',
            'created_at': 'Waktu penilaian diberikan.',
        },
        'catatan': 'Tabel dan relasinya sudah disiapkan di skema, namun fitur pemberian rating ini belum diimplementasikan di kode aplikasi/tampilan saat ini.',
    },
    'counselor_certifications': {
        'fungsi': 'Skema untuk menyimpan daftar sertifikasi profesional milik seorang konselor.',
        'kolom': {
            'certification_id': 'Kunci utama (PK).',
            'counselor_id': 'Konselor pemilik sertifikasi.',
            'certification_name': 'Nama sertifikasi.',
            'issuer': 'Lembaga penerbit.',
            'certificate_number': 'Nomor sertifikat.',
            'year': 'Tahun terbit.',
            'created_at': 'Waktu data dibuat.',
        },
        'catatan': 'Tabel dan relasinya sudah disiapkan di skema, namun belum ada menu/form di aplikasi untuk mengelola data ini.',
    },
    'diary_responses': {
        'fungsi': 'Skema untuk fitur balasan/tanggapan konselor terhadap sebuah catatan Diary yang dibagikan mahasiswa.',
        'kolom': {
            'id': 'Kunci utama (PK).',
            'diary_id': 'Catatan diary yang ditanggapi.',
            'counselor_id': 'Konselor yang memberi tanggapan.',
            'response': 'Isi tanggapan.',
            'created_at': 'Waktu tanggapan dibuat.',
        },
        'catatan': 'Tabel sudah disiapkan di skema, namun fitur balasan konselor atas Diary ini belum diimplementasikan — saat ini konselor hanya bisa melihat (read-only) Diary yang dibagikan lewat menu Diary Dibagikan.',
    },
    'notifications': {
        'fungsi': 'Skema untuk fitur notifikasi dalam-aplikasi per pengguna.',
        'kolom': {
            'id': 'Kunci utama (PK).',
            'user_id': 'Pemilik notifikasi.',
            'message': 'Isi pesan notifikasi.',
            'is_read': 'Penanda sudah/belum dibaca.',
            'created_at': 'Waktu notifikasi dibuat.',
        },
        'catatan': 'Belum digunakan oleh kode aplikasi saat ini — notifikasi antar-pengguna baru berjalan lewat email (lihat config/send_email.php), belum lewat tabel ini.',
    },
    'chat_logs': {
        'fungsi': 'Skema log chat generik dari versi awal, sebelum digantikan oleh chat_messages (yang membedakan pengirim/penerima secara eksplisit).',
        'kolom': {
            'id': 'Kunci utama (PK).',
            'user_id': 'Pengguna terkait log.',
            'sender': 'Penanda pengirim (mis. "user"/"bot").',
            'message': 'Isi pesan.',
        },
        'catatan': 'Tidak digunakan oleh kode aplikasi saat ini — sudah digantikan oleh chat_messages.',
    },
    'login_logs': {
        'fungsi': 'Skema untuk mencatat riwayat login pengguna (waktu & alamat IP).',
        'kolom': {
            'id': 'Kunci utama (PK).',
            'user_id': 'Pengguna yang login.',
            'login_time': 'Waktu login.',
            'ip_address': 'Alamat IP saat login.',
        },
        'catatan': 'Belum digunakan oleh kode aplikasi saat ini — proses login (AuthController) belum menulis ke tabel ini.',
    },
    'schema_migrations': {
        'fungsi': (
            'Bukan tabel data aplikasi, melainkan tabel infrastruktur milik skrip bin/migrate.php: mencatat nama '
            'file migrasi (database/migrations/*.sql) yang sudah pernah dijalankan pada database ini, agar skrip '
            'tidak menjalankan ulang migrasi yang sama saat aplikasi di-deploy ulang.'
        ),
        'kolom': {
            'migration': 'Kunci utama (PK) — nama file migrasi, mis. 2026_07_26_add_booking_cancellation_approval.sql.',
            'applied_at': 'Waktu migrasi tersebut dijalankan/ditandai selesai.',
        },
    },
}

KEY_LABEL = {'PRI': 'PK', 'UNI': 'Unik', 'MUL': 'FK/Indeks'}


def set_cell_shading(cell, color_hex):
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)
    cell._tc.get_or_add_tcPr().append(shd)


def set_col_widths(table, widths_cm):
    for row in table.rows:
        for cell, width in zip(row.cells, widths_cm):
            cell.width = Cm(width)


def add_heading(doc, text, level, color=NAVY, size=None):
    h = doc.add_heading(level=level)
    run = h.add_run(text)
    run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    return h


def add_table_relations(fks_by_table, table_name):
    rels = []
    for fk in fks_by_table.get(table_name, []):
        rels.append(f"{fk['column']} → {fk['ref_table']}.{fk['ref_column']}")
    return rels


def add_referenced_by(fks_by_ref_table, table_name):
    refs = []
    for fk in fks_by_ref_table.get(table_name, []):
        refs.append(f"{fk['table']}.{fk['column']} (mengacu ke {table_name}.{fk['ref_column']})")
    return refs


def main():
    out_path = sys.argv[1] if len(sys.argv) > 1 else os.path.join(ROOT, 'database', 'Dokumentasi_Struktur_Database_SIMKM.docx')

    env = load_env()
    print('Mengambil metadata skema dari database live...', file=sys.stderr)
    tables, columns, fks = fetch_schema(env)

    tables, columns, fks = merge_manual_tables(tables, columns, fks)

    fks_by_table = {}
    fks_by_ref_table = {}
    for fk in fks:
        fks_by_table.setdefault(fk['table'], []).append(fk)
        fks_by_ref_table.setdefault(fk['ref_table'], []).append(fk)

    documented = set()
    for _, table_list in MODULES:
        documented.update(table_list)
    missing = set(tables) - documented
    if missing:
        print(f'PERINGATAN: tabel berikut ada di DB tapi belum dikelompokkan di MODULES: {missing}', file=sys.stderr)

    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10.5)

    # --- Sampul ---
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.add_run('\n\n\n')
    run = title_p.add_run('DOKUMENTASI STRUKTUR BASIS DATA')
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = NAVY

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub_p.add_run('Sistem Informasi Manajemen Kesehatan Mental (SIMKM)')
    run.font.size = Pt(15)
    run.font.color.rgb = ACCENT

    info_p = doc.add_paragraph()
    info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info_p.add_run(f"Basis data: {env.get('DB_DATABASE', 'mental_health')}  |  Mesin: MySQL / MariaDB")
    run.font.size = Pt(11)
    run.font.color.rgb = MUTED

    doc.add_page_break()

    # --- Pendahuluan ---
    add_heading(doc, 'Pendahuluan', level=1)
    p = doc.add_paragraph(
        'Dokumen ini menjelaskan struktur basis data aplikasi SIMKM (Sistem Informasi Manajemen '
        'Kesehatan Mental) — aplikasi berbasis web yang membantu mahasiswa mencatat kondisi kesehatan '
        'mental melalui Diary Terstruktur, mengisi Self-Assessment (BDI-II & PWB), menggunakan fitur '
        'Self Help, serta berkonsultasi dengan konselor kampus melalui sistem booking dan chat. Admin '
        'mengelola akun, jadwal konselor, dan memantau melalui Laporan.'
    )
    p2 = doc.add_paragraph(
        'Basis data terdiri dari beberapa tabel yang dikelompokkan menurut modul fungsionalnya. Untuk '
        'setiap tabel dijelaskan: fungsi tabel secara umum, daftar kolom beserta kegunaannya, kunci '
        'utama (Primary Key/PK), kunci tamu (Foreign Key/FK) beserta tabel yang direferensikan, dan '
        'tabel lain yang balik mereferensikan tabel tersebut. Sebagian tabel merupakan peninggalan '
        '(legacy) dari iterasi awal pengembangan dan tidak lagi dipakai oleh kode aplikasi saat ini — '
        'hal ini disebutkan secara eksplisit pada bagian "Catatan" tabel yang bersangkutan, berdasarkan '
        'pemeriksaan langsung terhadap kode sumber (src/Controllers dan src/Repositories), bukan sekadar '
        'asumsi dari nama tabel/kolom.'
    )
    doc.add_paragraph(f'Total tabel yang didokumentasikan: {len(tables)}.')

    # --- Ringkasan modul ---
    doc.add_page_break()
    add_heading(doc, 'Ringkasan Modul & Tabel', level=1)
    doc.add_paragraph('Daftar modul aplikasi dan tabel-tabel basis data yang menjadi bagian dari modul tersebut.')

    summary_table = doc.add_table(rows=1, cols=2)
    summary_table.style = 'Light Grid Accent 1'
    summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = summary_table.rows[0].cells
    hdr[0].text = 'Modul'
    hdr[1].text = 'Tabel'
    for cell in hdr:
        for p_ in cell.paragraphs:
            for r in p_.runs:
                r.font.bold = True
    for module_name, table_list in MODULES:
        row = summary_table.add_row().cells
        row[0].text = module_name
        row[1].text = ', '.join(table_list)
    set_col_widths(summary_table, [5, 12])

    # --- Detail tabel ---
    doc.add_page_break()
    add_heading(doc, 'Detail Struktur Tabel', level=1)

    for module_name, table_list in MODULES:
        add_heading(doc, module_name, level=2, color=ACCENT)

        for table_name in table_list:
            if table_name not in columns:
                continue

            add_heading(doc, table_name, level=3)

            info = TABLE_DOCS.get(table_name, {})
            fungsi = info.get('fungsi', '')
            if fungsi:
                doc.add_paragraph(fungsi)

            catatan = info.get('catatan')
            if catatan:
                cat_p = doc.add_paragraph()
                run = cat_p.add_run('Catatan: ')
                run.bold = True
                run.font.color.rgb = RGBColor(0xB4, 0x53, 0x09)
                run2 = cat_p.add_run(catatan)
                run2.font.color.rgb = RGBColor(0xB4, 0x53, 0x09)
                run2.italic = True

            # Tabel kolom
            col_table = doc.add_table(rows=1, cols=3)
            col_table.style = 'Light Grid Accent 1'
            hdr = col_table.rows[0].cells
            hdr[0].text = 'Kolom'
            hdr[1].text = 'Tipe & Kunci'
            hdr[2].text = 'Keterangan'
            for cell in hdr:
                set_cell_shading(cell, '1E293B')
                for p_ in cell.paragraphs:
                    for r in p_.runs:
                        r.font.bold = True
                        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

            col_docs = info.get('kolom', {})
            for col in columns[table_name]:
                row = col_table.add_row().cells
                name_run = row[0].paragraphs[0].add_run(col['name'])
                name_run.font.bold = col['key'] == 'PRI'

                key_label = KEY_LABEL.get(col['key'], '')
                type_text = col['type']
                if key_label:
                    type_text += f" ({key_label})"
                if not col['nullable'] and col['key'] != 'PRI':
                    type_text += ' NOT NULL'
                row[1].text = type_text

                row[2].text = col_docs.get(col['name'], '-')

            set_col_widths(col_table, [3.5, 4.5, 9])

            # Relasi
            rels_out = add_table_relations(fks_by_table, table_name)
            rels_in = add_referenced_by(fks_by_ref_table, table_name)
            extra = info.get('relasi_extra', [])

            if rels_out or rels_in or extra:
                rel_p = doc.add_paragraph()
                rel_p.add_run('Relasi:').bold = True

                for r in rels_out:
                    doc.add_paragraph(f'Mengacu ke {r}', style='List Bullet')
                for r in rels_in:
                    doc.add_paragraph(f'Direferensikan oleh {r}', style='List Bullet')
                for r in extra:
                    doc.add_paragraph(r, style='List Bullet')
            else:
                doc.add_paragraph('Relasi: tidak memiliki foreign key ke/dari tabel lain.')

            doc.add_paragraph('')

    doc.save(out_path)
    print(f'Dokumen tersimpan di: {out_path}', file=sys.stderr)


if __name__ == '__main__':
    main()
